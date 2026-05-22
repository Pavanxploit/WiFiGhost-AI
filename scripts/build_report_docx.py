from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
REPORT_MD = DOCS / "project_report.md"
OUTPUT_DOCX = DOCS / "WiFiGhost_AI_Project_Report.docx"

ACCENT = RGBColor(11, 143, 138)
DARK = RGBColor(23, 33, 38)
MUTED = RGBColor(96, 110, 118)
LIGHT_FILL = "EAF5F3"
TABLE_HEADER = "DDEDEA"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "C8D6D8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "left")


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("WiFiGhost AI Project Report")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_header(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run("WiFiGhost AI - ML-Based Rogue Access Point Detection")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8D6D8")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_header(section)
    add_footer(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        style.font.color.rgb = ACCENT if style_name != "Heading 3" else DARK
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(5)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = ACCENT


def add_cover(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("WiFiGhost AI")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ML-Based Rogue Access Point Detection Using ESP32 and TFT Display")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("IoT + Machine Learning + Cybersecurity Project Report")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    document.add_paragraph()

    metadata = document.add_table(rows=5, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(metadata)
    rows = [
        ("Project Title", "WiFiGhost AI"),
        ("Domain", "Internet of Things, Machine Learning, Cybersecurity"),
        ("Hardware", "ESP32 DevKit, 2.8 inch SPI TFT Display, Jumper Wires, Breadboard"),
        ("Software", "Python, Flask, scikit-learn, Arduino IDE"),
        ("Prepared For", "Project Presentation and Submission"),
    ]
    for row, (key, value) in zip(metadata.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            if idx == 0:
                set_cell_fill(cell, LIGHT_FILL)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    document.add_paragraph()
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(
        "A low-cost defensive wireless security monitor that scans nearby access points, "
        "detects suspicious Wi-Fi behavior, and displays risk alerts on both web dashboard and TFT."
    )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = DARK
    document.add_page_break()


def add_image(
    document: Document,
    path: Path,
    caption: str,
    width_inches: float = 5.9,
    max_height_inches: float = 4.6,
) -> None:
    if not path.exists():
        return
    with Image.open(path) as img:
        width_px, height_px = img.size
    aspect = width_px / height_px
    target_width = width_inches
    target_height = target_width / aspect
    if target_height > max_height_inches:
        target_height = max_height_inches
        target_width = target_height * aspect

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(
        str(path),
        width=Inches(target_width),
        height=Inches(target_height),
    )
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED


def add_code_block(document: Document, code_lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_borders(cell, "D6E0E2")
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    set_cell_fill(cell, "F4F7F8")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Cascadia Mono"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Mono")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 52, 59)


def add_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            cell.text = value
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_fill(cell, TABLE_HEADER)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
                    if row_index == 0:
                        run.font.bold = True
    document.add_paragraph()


def add_paragraph_from_markdown(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("- "):
        paragraph = document.add_paragraph(stripped[2:], style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        return
    if re.match(r"^\d+\. ", stripped):
        paragraph = document.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tokens = re.split(r"(\*\*[^*]+\*\*)", stripped)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            paragraph.add_run(token)


def build_report() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    code_mode = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    for line in lines:
        if line.startswith("# "):
            continue

        if line.startswith("```"):
            if code_mode:
                add_code_block(document, code_lines)
                code_lines = []
                code_mode = False
            else:
                if table_lines:
                    add_table(document, table_lines)
                    table_lines = []
                code_mode = True
            continue

        if code_mode:
            code_lines.append(line)
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            add_table(document, table_lines)
            table_lines = []

        stripped = line.strip()
        if stripped.startswith("## "):
            text = stripped[3:]
            document.add_heading(text, level=1)
            if text == "5. Hardware and Software Requirements":
                add_image(document, DOCS / "assets" / "esp32-devkit.jpeg", "Figure 1: ESP32 DevKit used as the Wi-Fi scanner.", 3.1, 3.8)
                add_image(document, DOCS / "assets" / "tft-displays.jpeg", "Figure 2: Available TFT displays. The 2.8 inch SPI TFT is selected for ESP32 wiring.", 4.4, 3.6)
            if text == "13. Results":
                add_image(document, DOCS / "dashboard-alert.png", "Figure 3: Dashboard showing ALERT state for rogue AP demo.", 5.9)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        else:
            add_paragraph_from_markdown(document, line)

    if table_lines:
        add_table(document, table_lines)
    if code_lines:
        add_code_block(document, code_lines)

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
